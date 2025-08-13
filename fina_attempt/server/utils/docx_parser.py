# server/utils/docx_parser.py
from __future__ import annotations
import os
import re
import io
from typing import Dict, Iterator, List, Optional, Union

try:
    from docx import Document  # python-docx
except Exception as e:
    raise RuntimeError(
        "python-docx is required for DOCX parsing. Install with: pip install python-docx"
    ) from e

# Detect common requirement ID patterns seen in PRDs
REQ_ID_REGEX = re.compile(
    r"\b(?:Quasar[-_ ]?\d{3,6}|REQ[-_ ]?\d{2,6}|QSR[-_ ]?\d{2,6})\b",
    re.IGNORECASE,
)

NOISE_LINE = re.compile(r"^\s*(table of contents|revision history|confidential)\s*$", re.IGNORECASE)

def _clean(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s

def _is_noise(s: str) -> bool:
    if not s:
        return True
    if NOISE_LINE.match(s):
        return True
    if len(s) < 2:
        return True
    return False

def _iter_paragraph_text(doc: Document) -> Iterator[str]:
    for p in doc.paragraphs:
        t = _clean(p.text)
        if not _is_noise(t):
            yield t

def _iter_table_text(doc: Document) -> Iterator[str]:
    for tbl in doc.tables:
        for r in tbl.rows:
            for c in r.cells:
                t = _clean(c.text)
                if not _is_noise(t):
                    yield t

def extract_docx_blocks(source: Union[str, io.BytesIO], include_tables: bool = True) -> List[str]:
    """
    Return a flat list of text blocks from paragraphs (+ tables if enabled).
    Accepts a file path or a BytesIO stream.
    """
    if source is None:
        return []
    if isinstance(source, str):
        if not os.path.exists(source):
            return []
        doc = Document(source)
    else:
        source.seek(0)
        doc = Document(source)

    blocks: List[str] = list(_iter_paragraph_text(doc))
    if include_tables:
        blocks.extend(list(_iter_table_text(doc)))

    deduped: List[str] = []
    prev = None
    for b in blocks:
        if b != prev:
            deduped.append(b)
        prev = b
    return deduped

def detect_requirement_id(text: str) -> Optional[str]:
    m = REQ_ID_REGEX.search(text or "")
    return m.group(0) if m else None

def parse_docx_to_rows(
    source: Union[str, io.BytesIO],
    *,
    product: str = "unknown",
    subproduct: str = "unknown",
    file_label: Optional[str] = None,
    include_tables: bool = True,
    source_name: Optional[str] = None,  # backward compatibility
) -> List[Dict]:
    """
    Convert a DOCX PRD/spec into normalized rows.

    Accepts both the new and old calling styles.
    """
    # Backward-compat param mapping
    if source_name and not file_label:
        file_label = source_name

    file_name = file_label or (source if isinstance(source, str) else "uploaded.docx")
    if isinstance(file_name, str):
        file_name = os.path.basename(file_name)

    blocks = extract_docx_blocks(source, include_tables=include_tables)
    out: List[Dict] = []
    for i, t in enumerate(blocks):
        rid = detect_requirement_id(t)
        out.append(
            {
                "product": product,
                "subproduct": subproduct,
                "source_type": "prd",
                "file": file_name,
                "idx": i,
                "text": t,
                "requirement_id": rid,
            }
        )
    return out

